from __future__ import annotations

import json
import re
import shutil
import uuid
from pathlib import Path

from docx import Document
from pypdf import PdfReader

from .db import Database


ENTRY_RE = re.compile(
    r"^(?P<number>\*?\d{3,5}|M\d{3}|[αΩ]|\d{1,2}\s*[-–—]\s*\d{1,2})(?:\s*[:：|·\-–—]\s*(?P<title>.*))?$",
    re.IGNORECASE,
)


def normalize_text(value: str) -> str:
    value = value.replace("\r\n", "\n").replace("\r", "\n").replace("\u00a0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def extract_text(path: Path, page_start: int | None = None, page_end: int | None = None) -> str:
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return normalize_text(path.read_text(encoding="utf-8-sig"))
    if suffix == ".docx":
        doc = Document(path)
        lines = [paragraph.text for paragraph in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                lines.append("\t".join(cell.text for cell in row.cells))
        return normalize_text("\n".join(lines))
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        start = max(0, (page_start or 1) - 1)
        end = min(len(reader.pages), page_end or len(reader.pages))
        pages = [normalize_text(reader.pages[index].extract_text() or "") for index in range(start, end)]
        text = normalize_text("\n".join(pages))
        if len(re.sub(r"\s+", "", text)) < max(80, (end - start) * 20):
            raise ValueError("PDF 没有足够的文字层；当前版本不支持扫描图片 OCR")
        return text
    raise ValueError("仅支持 PDF、DOCX 和 TXT")


def split_segments(text: str, chapter_key: str, chapter_title: str) -> list[dict]:
    segments: list[dict] = []
    current: dict | None = None
    preface: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            if current:
                current["body_lines"].append("")
            elif preface:
                preface.append("")
            continue
        match = ENTRY_RE.match(stripped)
        if match:
            if current:
                segments.append(current)
            current = {
                "entry_number": re.sub(r"\s+", "", match.group("number")).replace("–", "-").replace("—", "-"),
                "title": (match.group("title") or "").strip(),
                "body_lines": [],
            }
        elif current:
            current["body_lines"].append(stripped)
        else:
            preface.append(stripped)
    if current:
        segments.append(current)
    if not segments:
        paragraphs = [part.strip() for part in re.split(r"\n\s*\n", text) if part.strip()]
        segments = [{"entry_number": str(index + 1), "title": "", "body_lines": [part]} for index, part in enumerate(paragraphs)]
    if preface:
        segments.insert(0, {"entry_number": "前言", "title": chapter_title, "body_lines": preface})
    output = []
    for order, segment in enumerate(segments):
        body = normalize_text("\n".join(segment.pop("body_lines")))
        if not body and not segment["title"]:
            continue
        output.append({**segment, "body": body, "chapter_key": chapter_key, "chapter_title": chapter_title, "sort_order": order, "reviewed": False})
    return output


def import_story(
    db: Database,
    library: Path,
    source_file: Path,
    book_id: str,
    title: str,
    chapter_key: str,
    chapter_title: str,
    page_start: int | None = None,
    page_end: int | None = None,
    source_name: str | None = None,
) -> dict:
    text = extract_text(source_file, page_start, page_end)
    segments = split_segments(text, chapter_key, chapter_title)
    if not segments:
        raise ValueError("没有提取到可导入的故事段落")
    source_dir = library / "sources" / "stories"
    source_dir.mkdir(parents=True, exist_ok=True)
    stored = source_dir / f"{uuid.uuid4().hex}{source_file.suffix.lower()}"
    shutil.copy2(source_file, stored)
    with db.connect() as conn:
        conn.execute(
            """INSERT INTO story_books(id,title,source_name,source_path,status)
            VALUES(?,?,?,?, 'review') ON CONFLICT(id) DO UPDATE SET title=excluded.title,
            source_name=excluded.source_name,source_path=excluded.source_path,status='review',updated_at=CURRENT_TIMESTAMP""",
            (book_id, title, source_name or source_file.name, stored.relative_to(library).as_posix()),
        )
        conn.execute("DELETE FROM story_segments WHERE book_id=? AND chapter_key=?", (book_id, chapter_key))
        conn.executemany(
            """INSERT INTO story_segments
            (book_id,chapter_key,chapter_title,entry_number,title,body,sort_order,reviewed)
            VALUES(?,?,?,?,?,?,?,?)""",
            [(book_id, s["chapter_key"], s["chapter_title"], s["entry_number"], s["title"], s["body"], s["sort_order"], 0) for s in segments],
        )
    return {"book_id": book_id, "segments": len(segments), "status": "review"}


def story_books(db: Database) -> list[dict]:
    return db.all("""
      SELECT b.*,COUNT(s.id) AS segment_count,SUM(s.reviewed) AS reviewed_count
      FROM story_books b LEFT JOIN story_segments s ON s.book_id=b.id
      GROUP BY b.id ORDER BY b.id
    """)


def story_segments(db: Database, book_id: str) -> list[dict]:
    return db.all("SELECT * FROM story_segments WHERE book_id=? ORDER BY chapter_key,sort_order,id", (book_id,))


def update_segment(db: Database, segment_id: int, payload: dict) -> None:
    row = db.one("SELECT * FROM story_segments WHERE id=?", (segment_id,))
    if not row:
        raise ValueError("故事段落不存在")
    db.execute(
        """UPDATE story_segments SET entry_number=?,title=?,body=?,chapter_key=?,chapter_title=?,sort_order=?,reviewed=? WHERE id=?""",
        (
            str(payload.get("entry_number", row["entry_number"])), str(payload.get("title", row["title"])),
            str(payload.get("body", row["body"])), str(payload.get("chapter_key", row["chapter_key"])),
            str(payload.get("chapter_title", row["chapter_title"])), int(payload.get("sort_order", row["sort_order"])),
            int(bool(payload.get("reviewed", row["reviewed"]))), segment_id,
        ),
    )


def split_segment(db: Database, segment_id: int, offset: int) -> int:
    row = db.one("SELECT * FROM story_segments WHERE id=?", (segment_id,))
    if not row:
        raise ValueError("故事段落不存在")
    body = row["body"]
    if offset <= 0 or offset >= len(body):
        raise ValueError("请选择正文中间的位置进行拆分")
    left, right = body[:offset].rstrip(), body[offset:].lstrip()
    if not left or not right:
        raise ValueError("拆分后的两段都必须有正文")
    with db.connect() as conn:
        conn.execute(
            "UPDATE story_segments SET sort_order=sort_order+1 WHERE book_id=? AND chapter_key=? AND sort_order>?",
            (row["book_id"], row["chapter_key"], row["sort_order"]),
        )
        conn.execute("UPDATE story_segments SET body=?,reviewed=0 WHERE id=?", (left, segment_id))
        cursor = conn.execute(
            """INSERT INTO story_segments(book_id,chapter_key,chapter_title,entry_number,title,body,sort_order,reviewed)
            VALUES(?,?,?,?,?,?,?,0)""",
            (row["book_id"], row["chapter_key"], row["chapter_title"], f"{row['entry_number']}-2", "", right, row["sort_order"] + 1),
        )
        return int(cursor.lastrowid)


def merge_next_segment(db: Database, segment_id: int) -> int:
    row = db.one("SELECT * FROM story_segments WHERE id=?", (segment_id,))
    if not row:
        raise ValueError("故事段落不存在")
    next_row = db.one(
        """SELECT * FROM story_segments WHERE book_id=? AND chapter_key=? AND sort_order>? ORDER BY sort_order,id LIMIT 1""",
        (row["book_id"], row["chapter_key"], row["sort_order"]),
    )
    if not next_row:
        raise ValueError("当前章节没有下一段可合并")
    merged = normalize_text("\n\n".join(filter(None, (row["body"], next_row["body"]))))
    with db.connect() as conn:
        conn.execute("UPDATE story_segments SET body=?,reviewed=0 WHERE id=?", (merged, segment_id))
        conn.execute("DELETE FROM story_segments WHERE id=?", (next_row["id"],))
        conn.execute(
            "UPDATE story_segments SET sort_order=sort_order-1 WHERE book_id=? AND chapter_key=? AND sort_order>?",
            (row["book_id"], row["chapter_key"], next_row["sort_order"]),
        )
    return int(next_row["id"])


def storybook_payload(
    db: Database, reviewed_only: bool = False, book_ids: set[str] | None = None,
    omit_empty: bool = False,
) -> dict:
    books = []
    for book in story_books(db):
        if book_ids is not None and book["id"] not in book_ids:
            continue
        segments = story_segments(db, book["id"])
        if reviewed_only:
            segments = [segment for segment in segments if segment["reviewed"]]
        if omit_empty and not segments:
            continue
        chapters = []
        seen = set()
        entries = []
        for segment in segments:
            if segment["chapter_key"] not in seen:
                chapters.append({"key": segment["chapter_key"], "title": segment["chapter_title"]})
                seen.add(segment["chapter_key"])
            try:
                metadata = json.loads(segment.get("metadata_json") or "{}")
            except (TypeError, json.JSONDecodeError):
                metadata = {}
            if not isinstance(metadata, dict):
                metadata = {}
            entry = {
                **metadata,
                "key": metadata.get("key") or f"{book['id']}-{segment['chapter_key']}-{segment['id']}",
                "id": segment["entry_number"], "title": segment["title"],
                "chapterKey": segment["chapter_key"], "chapter": segment["chapter_title"],
                "text": segment["body"],
            }
            entry.setdefault("section", entry.get("encounter") or segment["chapter_title"])
            entries.append(entry)
        books.append({"id": book["id"], "title": book["title"], "entryCount": len(entries), "chapters": chapters, "entries": entries})
    return {"generatedAt": "ATO Asset Studio", "books": books}


def storybook_javascript(
    db: Database, reviewed_only: bool = True, book_ids: set[str] | None = None,
) -> bytes:
    payload = json.dumps(
        storybook_payload(
            db, reviewed_only=reviewed_only, book_ids=book_ids, omit_empty=reviewed_only,
        ),
        ensure_ascii=False, separators=(",", ":"),
    )
    return f"window.STORYBOOK_DATA = {payload};\n".encode("utf-8")
